#import google.generativeai as genai
from celery import shared_task
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import ast
import re
import os
from typing import Dict, Any
import google.generativeai as genai


# Configura a API Key (nunca exponha publicamente)
#genai.configure(api_key="AIzaSyCYfHeJAJqyi2qrBpygTQn0XxTfg4K83po")
genai.configure(api_key="AIzaSyBvpQ7uT06MgZd72_9ba6mCEbSUk8iosJI")

# ---------- Configuração do Gemini ----------
#genai.configure(api_key="AIzaSyDdLN5FlE_mJdsH2yQcSWCBJcMsceJOLWs")

modelo = genai.GenerativeModel("gemini-2.5-flash")


# ---------- Helpers ----------
def adicionar_numero_paginas(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def limpar_texto(texto: str) -> str:
    texto = re.sub(r"[*#_~>`-]+", "", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def sugerir_ideias(dados: str, pergunta: str, qtd_caracteres: str):
    prompt = f"""
    Gere *exatamente* 5 respostas curtas com no máximo {qtd_caracteres} caracteres
    para a pergunta: "{pergunta}".

    BASE: {dados}

    REGRAS OBRIGATÓRIAS:
    - Responda somente com uma lista Python válida de 5 strings.
    - Nenhuma string pode ultrapassar {qtd_caracteres} caracteres.
    - Não inclua explicações, títulos, código, prefixos ("python") ou texto fora da lista.
    - Não escreva frases longas.
    - Não adicione comentários.

    Exemplo de formato correto:
    ["ideia 1", "ideia 2", "ideia 3", "ideia 4", "ideia 5"]
    """

    resposta = modelo.generate_content(prompt)
    texto = limpar_texto(resposta.text)

    # 🔥 transforma texto "[...]" em lista Python real
    try:
        lista = ast.literal_eval(texto)
        if isinstance(lista, list):
            return lista
        else:
            return []
    except:
        return []


# ---------- Função principal ----------
#@shared_task
@shared_task(time_limit=60, soft_time_limit=50)
def gerar_conteudo_secoes(dados: str) -> dict[str, Any]:

    prompt = f"""
    Gere um PLANO DE NEGÓCIO COMPLETO, claro e profissional.

    DADOS DA EMPRESA:
    {dados}

    REGRAS OBRIGATÓRIAS:
    - Cada secção deve ter no máximo 30 frases curtas.
    - Linguagem simples, objetiva e prática.
    - Não use markdown.
    - Não use listas com bullets.
    - Responda APENAS com um JSON válido exatamente neste formato:

    {{
      "SUMARIO_EXECUTIVO": "...",
      "A_EMPRESA": "...",
      "CARATERIZACAO_DO_PROJETO": "...",
      "O_PRODUTO_SERVICO": "...",
      "ANALISE_DE_MERCADO": "...",
      "SISTEMA_PRODUTIVO": "...",
      "PLANO_DE_MARKETING": "...",
      "ESTRUTURA_ORGANIZACIONAL": "...",
      "PLANO_FINANCEIRO": "...",
      "CONSIDERACOES_FINAIS": "...",
      "ANEXOS": "..."
    }}
    """

    resposta = modelo.generate_content(prompt)
    texto = limpar_texto(resposta.text)

    try:
        resultado = ast.literal_eval(texto)
        if isinstance(resultado, dict):
            return resultado
        return {}
    except Exception as e:
        return {"erro": "Falha ao gerar plano"}



def gerar_plano_de_negocio_word(tipo_negocio: str, localizacao: str, imagem_capa="capa.jpg"):
    doc = Document()

    # CAPA
    doc.add_paragraph("\n\n")

    # Inserir imagem de capa se existir
    if os.path.exists(imagem_capa):
        doc.add_picture(imagem_capa, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("\n")
    titulo = doc.add_paragraph("PLANO DE NEGÓCIO")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titulo.runs[0]
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph("\n")
    nome = doc.add_paragraph(f"Tipo de negócio: {tipo_negocio}")
    nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nome.runs[0].font.size = Pt(20)
    nome.runs[0].bold = True

    doc.add_paragraph("\n")
    local = doc.add_paragraph(f"Localização: {localizacao}")
    local.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    local.runs[0].font.size = Pt(16)

    doc.add_paragraph("\n\n")
    autor = doc.add_paragraph("Elaborado por: Adelino Francisco Emiliano")
    autor.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    autor.runs[0].italic = True

    data = doc.add_paragraph("Data: 2025")
    data.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    data.runs[0].italic = True

    doc.add_page_break()

    # Adicionar número de páginas (exceto capa)
    for section in doc.sections[1:]:
        adicionar_numero_paginas(section)

    # Gerar conteúdo detalhado automaticamente
    dados = {
        "tipo_negocio": tipo_negocio,
        "localizacao": localizacao
    }
    conteudo_secoes = gerar_conteudo_secoes(dados)

    # Inserir conteúdo no documento Word
    for titulo_secao, texto in conteudo_secoes.items():
        doc.add_heading(titulo_secao, level=1)
        for paragrafo in texto.split("\n"):
            if paragrafo.strip():
                doc.add_paragraph(paragrafo.strip())
        doc.add_page_break()

    doc.save("plano_de_negocio.docx")
    print("📄 Plano de negócio Word gerado com sucesso: plano_de_negocio.docx")


#gerar_plano_de_negocio_word(dados_exemplo, imagem_capa="capa.jpg")
